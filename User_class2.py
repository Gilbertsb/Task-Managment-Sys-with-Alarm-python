# AUTHOR:GILBERT
# this class is a child class of the child class Main_task which means we did Multilevel inheritance
# this is the user class allows users to perform the different operations we
# we have in the Main_task and Sub_task
# contains methods, control flow, menu of the actions the user
# might want to perform

import time
from winsound import Beep
from main_task_class import *
from sub_task_class import *
import colorama
from colorama import Fore, Back, Style
import docx

colorama.init()

# created empty list and strings and assign to variables
# main_task = main_task we did not want to confuse the system
now = ''
my_date = ''
get_me = ""
my_date_dc = {}
sub_name = {}
sub_location = {}
sub_start = {}
sub_task_start=''
main_tas = []
task_time_start = ''
task_time_end = ''
doc2 = docx.Document()
user_name_n = ''


class User_in_main(Main_task):
    # define init function that takes in arguments

    def __init__(self, my_date, task_name, location, time_start, time_end):

        super().__init__(my_date, task_name, location, time_start, time_end)

    @staticmethod
    # this method allows user to add their name
    def user_name():
        global user_name_n
        print("+------------------------------------+")
        print("| We want to keep track of           |")
        print("| what your are doing in our system  |")
        print("| your names please                  |")
        print("+------------------------------------+")

        user_name_n = str(input("ENTER YOUR NAME PLEASE:"))

        if user_name_n.isalpha() == False:
            print(Fore.RED + 'USE CHARACTERS ONLY' + Fore.WHITE)
            User_in_main.user_name()

    # @stactic makes method in every class accessible to this file
    @staticmethod
    # this method allows users to add task
    def add_task():
        global user_name_n
        global get_me
        global main_tas
        global my_date
        global task_time_start
        global task_time_end

        from datetime import datetime
        now = datetime.now()
        print("+---------------------------+")
        print('|1.SET FOR CURRENT DAY     |')
        print('|2.SET FOR DIFFERENT DAY   |')
        print("+---------------------------+")
        try:#validation
            choice = input('CHOOSE A WAY TO SET DATE: ')
            if choice == '1':
                my_date = now.date()
            elif choice == '2':
                user_date = input('Set Date for you task(YYYY-MM-DD format): ')
                date_time_obj = datetime.strptime(user_date, '%Y-%m-%d')
                my_date = date_time_obj.date()
            else:
                print("YOU DIDN'T CHOOSE ANY THING!!")
                User_in_main.add_task()
        except ValueError:#what to show if user enter incorrect values
            print(Fore.RED + 'RE-RUN THE PROGRAM AND USE THE RIGHT FORMAT PLEASE!! (YYYY-MM-DD format)' + Fore.WHITE)
            exit()

        task_name = input("Enter main Task name(ex: Teaching): ")
        if task_name.isalpha() == False:
            print(Fore.RED + 'USE CHARACTERS PLEASE' + Fore.WHITE)
            User_in_main.add_task()
        task_location = input("Enter main Task location(ex: KH: )")
        if task_location.isalpha() == False:
            print(Fore.RED + 'USE CHARACTERS PLEASE' + Fore.WHITE)
            User_in_main.add_task()
        try:#validation
            task_time_star = input("Enter the expected time to start your task(ex: 12:03:10 : )")
            date_time_obj = datetime.strptime(task_time_star, '%H:%M:%S')
            task_time_sta = date_time_obj.time()
            task_time_start=str(task_time_sta)
        except ValueError:#what to show if user enter incorrect values
            print(Fore.RED + 'USE THE RIGHT FORMAT(HH:MM:SS)' + Fore.WHITE)
            User_in_main.add_task()


        try:#validation
            task_time_en = input("Enter the expected time to end your task(ex: 12:03:10 : )")
            date_time_obj = datetime.strptime(task_time_en, '%H:%M:%S')
            task_time_e = date_time_obj.time()
            task_time_end=str(task_time_e)
        except ValueError:#what to show if user enter incorrect values
            print(Fore.RED + 'USE THE RIGHT FORMAT(HH:MM:SS)' + Fore.WHITE)
            User_in_main.add_task()

        # add users inputs to the main task list
        main_tas.append(Main_task(my_date, task_name, task_location, task_time_start, task_time_end))
        get_me = task_name
        # returns the number corresponding to the main task user enters
        print(Fore.CYAN + "YOUR TASK HAS BEEN ADDED SUCCESSFULLY TO THE LIST!!" + Fore.WHITE)
        doc2.add_heading(user_name_n + " has Chosen 1 from menu and that is"
                                       " ADD NEW TASK menu see what they did: ")

        # this records whatever the user is doing in our system
        doc2.add_paragraph(" ______________________________________________________________")
        doc2.add_paragraph("Entered their Date which is : " + str(my_date))
        doc2.add_paragraph("Entered their task name which is : " + task_name)
        doc2.add_paragraph("Entered their location for the task which is: " + task_location)
        doc2.add_paragraph("Entered their time to start task which is : " + str(task_time_start))
        doc2.add_paragraph("Entered their time to End task which is: " + str(task_time_end))
        doc2.save(user_name_n + '.docx')

    @staticmethod
    # this method allow users to view their main_task
    def view_task_m():
        global user_name_n
        doc2.add_heading(user_name_n + " has visited chosen 2 from menu and that is"
                                       " VIEW TASK menu: ")
        doc2.add_paragraph(" ___________________________________________")
        doc2.add_paragraph("And they view MAIN TASKS")
        doc2.save(user_name_n + '.docx')

        print("Your main tasks are:")
        print(" TASK NAME      TASK LOCATION      TIME-START      TIME-END")
        print("--------------------------------------------------------------")
        # created a for loop function so its print out the info in that order
        for i in main_tas:
            print(i.task_name, '--', i.location, '--', i.time_start, '--', i.time_end)
        print("--------------------------------------------------------------")


# this is the user class allows users to perform the different operations we
# we have in the sub_task
class User_in_sub(Sub_task):
    global get_me

    # define init function that takes in arguments
    def __init__(self, my_date, task_name, location, time_start_sub):
        super().__init__(my_date, task_name, location, time_start_sub)

    @staticmethod
    # after user has added main task,
    # this method ask user if they want to add a subtask eg
    # when you have enter main_task or to do Teaching then whether they are teaching
    # cohort 1 and cohort 2
    def add_new_sub():
        global user_name_n
        global sub_name
        global sub_location
        global sub_start
        global my_date
        global my_date_dc
        global sub_task_start
        from datetime import datetime

        # the prompts that ask user if they want to add a subtask
        sub_task_prompt = input("Do you have any sub task under this main task? \n "
                                "like Main: Teaching and Sub: Teaching cohort 1 Enter(Y/N): ")
        # control flow to prints out based on user input,
        # the case.fold takes in the user input in
        # whatever form and the code can still work
        # correctly
        if sub_task_prompt.casefold() == "Y".casefold():

            print("-------------------------------------------")
            sub_task_name = input("Enter name of sub task(ex Teaching cohort 1):")
            if sub_task_name.isalpha() == False:#validation
                print(Fore.RED + 'USE CHARACTERS PLEASE' + Fore.WHITE)
                User_in_sub.add_new_sub()
            sub_task_location = input("Enter location of sub task (ex Mali):")
            if sub_task_location.isalpha() == False:#validation
                print(Fore.RED + 'USE CHARACTERS PLEASE' + Fore.WHITE)
                User_in_sub.add_new_sub()
            try:#validation
                sub_task_sta = input("Enter the expected time to start your sub task(ex: 12:03:10 : )")
                date_time_obj = datetime.strptime(sub_task_sta, '%H:%M:%S')
                sub_task_star= date_time_obj.time()
                sub_task_start=str(sub_task_star)
            except ValueError:#what to show if user enter incorrect values
                print(Fore.RED + 'USE THE RIGHT FORMAT(HH:MM:SS)' + Fore.WHITE)
                User_in_sub.add_new_sub()

            # created variables that will take in the user_input
            # for the sub_task entered
            my_date_dc[get_me] = my_date
            sub_name[get_me] = sub_task_name
            sub_location[get_me] = sub_task_location
            sub_start[get_me] = sub_task_start

            # this function records everything the usser does on our system
            doc2.add_paragraph("User choose Y as they want to add sub task under the main task"
                               "they entered and the Sub task name is : " + sub_task_name)
            doc2.add_paragraph("Entered Sub task location which : " + sub_task_location)
            doc2.add_paragraph("Entered Expected time to start sub task which is : " + sub_task_start)
            doc2.save(user_name_n + '.docx')

        elif sub_task_prompt.casefold() == "N".casefold():
            print("Your task: " + str(get_me) + " has been successfully added!")
        # -----------------------------------

    @staticmethod
    # this method allows users to view their sub_task
    def view_task():
        global user_name_n

        doc2.add_heading(user_name_n + " has chosen 2 from the main menu and that is"
                                       " VIEW TASK menu: ")
        doc2.add_paragraph(" ___________________________________________")
        doc2.add_paragraph("And the user also viewed their SUB TASKS")
        doc2.save(user_name_n + '.docx')

        for i in sub_name:
            print("+------------------------------------------+")
            print('Main task is:', i.upper(), '            ')
            print("+------------------------------------------+")
            print("Your sub tasks are:")
            print('\t', 'Name: ', sub_name[i], "            ")
            print('\t', 'Location: ', sub_location[i], "    ")
            print('\t', 'Time: ', sub_start[i], "           ")
            print("+------------------------------------------+")

    @staticmethod
    # this method allows users to remove_task main task
    # if they feel like it has been completed
    def remove_task_m():
        global user_name_n
        global main_tas
        global sub_name
        global sub_location
        global sub_start
        remov=''

        # records what the user is doing in our system
        doc2.add_heading(user_name_n + " has chosen 3 from the menu"
                                       " and that is REMOVE A TASK see what they did: ")
        doc2.add_paragraph(" ___________________________________________")
        doc2.add_paragraph(" user removed their task : ")

        index = 1
        for i in main_tas:
            print(index, '.', i.task_name)
            index += 1


        try:#validation
            remov = input("Select Task(ex:1) you want to remove: ")
        except:#what to show if user enter incorrect values
            if remov.isdigit()==False:
                print(Fore.RED + "THAT CAN NOT WORK!" + Fore.WHITE)

        try:#validation
            task_rem = main_tas[int(remov) - 1]
            for c in main_tas:  # loop to check all main tasks
                if task_rem == c:  # checking if selected tasks is in our list
                    doc2.add_paragraph(c.task_name)
                    doc2.add_paragraph("-----------")
                    doc2.save(user_name_n + '.docx')
                    main_tas.remove(task_rem)
                    print("YOUR TASK HAS SUCCESSFULLY REMOVED FROM THE LIST!!")
                    for f in list(sub_name.keys()):
                        if task_rem == f:
                            del my_date_dc[f]
                            del sub_name[f]
                            del sub_location[f]
                            del sub_start[f]
                            print('Your task has been removed')
        except ValueError:#what to show if user enter incorrect values
            print(Fore.RED+"THAT CAN NOT WORK!"+Fore.WHITE)
        except IndexError:#what to show if user enter incorrect values
            print(Fore.RED+"IT IS NOT ON THE LIST!"+Fore.WHITE)


    @staticmethod
    # this method allows user to start there task and
    # start the time
    def start_time():
        global user_name_n
        global main_tas
        global sub_name
        global sub_location
        global sub_start
        global now
        global my_date_dc

        from datetime import datetime
        now = datetime.now()
        print('Current time: ', now.time().strftime("%H:%M:%S"))
        print("-------------------------------")
        print("Your tasks start now")
        doc2.add_heading(user_name_n + " has visited START YOUR TIMER menu see what he did: ")
        doc2.add_paragraph(" ___________________________________________")
        doc2.add_paragraph("And Current time was :  " + now.time().strftime("%H:%M:%S"))
        doc2.save(user_name_n + '.docx')

        while True:
            from datetime import datetime
            now = datetime.now()
            my_current_dt = now.date()
            y = datetime.now().strftime("%H:%M:%S")
            print(y)
            time.sleep(1)
            for x in main_tas:
                if x.time_start == y and x.my_date == my_current_dt:
                    freq = 1500
                    duration = 200
                    print(y, Fore.BLUE + '(You are doing:,', x.task_name, ')')  # changig color
                    for i in range(1, 30):
                        Beep(freq, duration)
            for c in sub_start:
                if sub_start[c] == y and my_date_dc[c] == my_current_dt:
                    freq = 1500
                    duration = 200
                    print(y, Fore.GREEN + '(You are doing:,', sub_name[c], ')')  # changig color
                    for i in range(1, 7):
                        Beep(freq, duration)

    @staticmethod
    # this methods allows users to see the menu and use ot
    def main_menu():
        global now
        from datetime import datetime
        now = datetime.now()
        print("WELCOME TO YOUR TODO MENU! BELOW ARE WHAT WE OFFER")
        print(Fore.GREEN + 'Current Date: ', now.date(), Fore.WHITE)
        print(Fore.YELLOW + 'Current time: ', now.time().strftime("%H:%M:%S") + Fore.WHITE)
        print("__________________________________________")
        print("|\t 1.ADD NEW TASK                       |")
        print("|\t 2.VIEW YOUR TASK                     |")
        print("|\t 3.REMOVE A TASK                      |")
        print("|\t 4.START YOUR TIMER                   |")
        print("|_________________________________________|")
        # prompting user to select one of the menu
        selection = input("Select from menu with using number(1-4): ")

        if selection == '1':  # checking if user selected one
            User_in_main.add_task()  # calling function to add task
            User_in_sub.add_new_sub()
        elif selection == '2':  # checking if user selected two
            print("+---------------------------+")
            print('|1.MAIN TASK                |')
            print('|2.MAIN AND THEIR SUB TASKS |')
            print("+---------------------------+")
            choose =input('CHOOSE TYPE OF TASK TO VIEW: ')
            if choose == '1':
                User_in_main.view_task_m()  # calling function to view task
            elif choose == '2':
                User_in_sub.view_task()
            else:
                print(Fore.RED + "YOU DIDN'T CHOOSE ANYTHING" + Fore.WHITE)


        elif selection == '3':  # checking if user selected three
            User_in_sub.remove_task_m()  # calling function to remove tasks
        elif selection == '4':  # checking if user selected four
            User_in_sub.start_time()  # calling function to start timer
        else:
            # if user doesn't  choose any of above numbers
            print(Fore.RED + "INVALID INPUT PLEASE TRY AGAIN!!"+ Fore.WHITE)

    @staticmethod
    # allows user to exit the program
    def exit_this():
        global user_name_n
        doc2.add_paragraph(user_name_n + " has chosen to exit the app: ")
        doc2.add_paragraph(" ___________________________________________")
        doc2.save(user_name_n + '.docx')
        exit(Fore.CYAN +"THANK YOU FOR USING THIS TODO SOFTWARE SEE YOUUU!!!"+ Fore.WHITE)
